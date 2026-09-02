"""Geração do documento de Ordem de Serviço em DOCX.

`montar_dados_os()` é a função única de montagem de dados — também consumida
por `service_order_pdf.py`, para os dois formatos nunca divergirem sobre o
que aparece no documento (só a forma de desenhar muda).

Dados da MasterSat aqui são uma cópia deliberada dos mesmos usados em
`contract_pdf.py`/`boleto_pdf.py` (ver comentário lá) — não foi criado um
módulo central de dados da empresa nesta fase para não alterar esses dois
arquivos existentes; fica registrado como débito técnico para uma limpeza
futura.
"""
from __future__ import annotations

import base64
import io
import logging
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.models.client import Client
from app.models.document import Document
from app.models.service_order import ServiceOrder
from app.models.service_order_material import ServiceOrderMaterial
from app.models.service_order_status_log import ServiceOrderStatusLog
from app.models.tracker import Tracker
from app.models.user import User
from app.models.vehicle import Vehicle
from app.services.storage import get_object_stream

logger = logging.getLogger(__name__)

# ── Dados da MasterSat (mesma empresa de contract_pdf.py, cópia deliberada) ──
EMPRESA_RAZAO = 'MASTERSAT COMERCIO E SERVIÇOS DE RASTREAMENTO LTDA'
EMPRESA_CNPJ = '14.228.344/0001-67'
EMPRESA_ENDERECO = 'RUA MARITIMA, 424 COMASA JOINVILLE SC'

_LOGO_PATH = Path(__file__).parent.parent / 'static' / 'mastersat_logo.png'
_MAX_EMBEDDED_PHOTOS = 12

TITULOS = {
    'ordem_servico': 'Ordem de Serviço',
    'termo_instalacao': 'Termo de Instalação',
    'termo_retirada': 'Termo de Retirada de Equipamento',
    'historico_execucao': 'Histórico de Execução',
}


@dataclass
class ChecklistItemData:
    description: str
    done: bool
    notes: str | None = None


@dataclass
class MaterialData:
    description: str
    quantity: str
    unit: str | None
    unit_price: str | None


@dataclass
class PhotoData:
    file_name: str
    content: bytes


@dataclass
class StatusLogData:
    when: str
    previous: str
    new: str
    notes: str | None


@dataclass
class OSDocumentData:
    kind: str
    order: ServiceOrder
    client: Client | None
    vehicle: Vehicle | None
    tracker: Tracker | None
    technician: User | None
    checklist: list[ChecklistItemData] = field(default_factory=list)
    materials: list[MaterialData] = field(default_factory=list)
    photos: list[PhotoData] = field(default_factory=list)
    technician_signature_png: bytes | None = None
    client_signature_png: bytes | None = None
    status_logs: list[StatusLogData] | None = None
    generated_at: datetime = field(default_factory=datetime.utcnow)

    @property
    def title(self) -> str:
        return TITULOS.get(self.kind, 'Documento Operacional')


def _read_object_bytes(object_key: str) -> bytes | None:
    try:
        obj = get_object_stream(object_key)
    except Exception:
        logger.warning('Falha ao baixar objeto %s do storage para documento de OS', object_key, exc_info=True)
        return None
    try:
        return obj.read()
    finally:
        try:
            obj.close()
            obj.release_conn()
        except Exception:  # noqa: BLE001 — limpeza de conexão não pode quebrar a geração do documento
            pass


def _checklist_items(raw: dict | list | None) -> list[ChecklistItemData]:
    """Mesma tolerância de formato do schema Pydantic (`_coerce_checklist`),
    aplicada aqui direto sobre o valor cru da coluna JSON."""
    if raw is None:
        return []
    items = raw.get('items') if isinstance(raw, dict) else raw
    if not isinstance(items, list):
        return []
    result = []
    for entry in items:
        if isinstance(entry, str):
            result.append(ChecklistItemData(description=entry, done=False, notes=None))
        elif isinstance(entry, dict):
            result.append(ChecklistItemData(
                description=str(entry.get('description', '')),
                done=bool(entry.get('done', False)),
                notes=entry.get('notes'),
            ))
    return result


def _load_photos(db: Session, order_id: int) -> list[PhotoData]:
    docs = db.scalars(
        select(Document)
        .where(
            Document.reference_type == 'service_order',
            Document.reference_id == order_id,
            Document.active.is_(True),
            Document.content_type.in_(('image/jpeg', 'image/png', 'image/webp')),
        )
        .order_by(Document.id.asc())
        .limit(_MAX_EMBEDDED_PHOTOS)
    ).all()
    photos: list[PhotoData] = []
    for doc in docs:
        content = _read_object_bytes(doc.object_key)
        if content:
            photos.append(PhotoData(file_name=doc.file_name, content=content))
    return photos


def _load_signature(db: Session, document_id: int | None) -> bytes | None:
    if not document_id:
        return None
    doc = db.get(Document, document_id)
    if not doc or not doc.active:
        return None
    return _read_object_bytes(doc.object_key)


def montar_dados_os(kind: str, order: ServiceOrder, db: Session) -> OSDocumentData:
    """Fonte única de dados para os dois geradores (DOCX e PDF)."""
    client = db.get(Client, order.client_id) if order.client_id else None
    vehicle = db.get(Vehicle, order.vehicle_id) if order.vehicle_id else None
    tracker = db.get(Tracker, order.tracker_id) if order.tracker_id else None
    technician = db.get(User, order.technician_id) if order.technician_id else None

    material_rows = db.scalars(
        select(ServiceOrderMaterial)
        .where(
            ServiceOrderMaterial.service_order_id == order.id,
            ServiceOrderMaterial.is_deleted.is_(False),
        )
        .order_by(ServiceOrderMaterial.id.asc())
    ).all()
    materials = [
        MaterialData(
            description=m.description,
            quantity=f'{m.quantity:g}' if m.quantity is not None else '-',
            unit=m.unit,
            unit_price=f'R$ {m.unit_price:.2f}' if m.unit_price is not None else None,
        )
        for m in material_rows
    ]

    status_logs = None
    if kind == 'historico_execucao':
        rows = db.scalars(
            select(ServiceOrderStatusLog)
            .where(ServiceOrderStatusLog.service_order_id == order.id)
            .order_by(ServiceOrderStatusLog.created_at.asc())
        ).all()
        status_logs = [
            StatusLogData(
                when=entry.created_at.strftime('%d/%m/%Y %H:%M') if entry.created_at else '-',
                previous=entry.previous_status.value if entry.previous_status else 'novo',
                new=entry.new_status.value,
                notes=entry.notes,
            )
            for entry in rows
        ]

    return OSDocumentData(
        kind=kind,
        order=order,
        client=client,
        vehicle=vehicle,
        tracker=tracker,
        technician=technician,
        checklist=_checklist_items(order.checklist),
        materials=materials,
        photos=_load_photos(db, order.id),
        technician_signature_png=_load_signature(db, order.technician_signature_document_id),
        client_signature_png=_load_signature(db, order.client_signature_document_id),
        status_logs=status_logs,
    )


def decode_signature_image(image_base64: str) -> bytes:
    """Decodifica o PNG capturado no canvas de assinatura do frontend —
    aceita tanto data URL (``data:image/png;base64,...``) quanto base64 puro."""
    payload = image_base64.split(',', 1)[1] if ',' in image_base64 and image_base64.strip().startswith('data:') else image_base64
    return base64.b64decode(payload)


# ── Helpers de layout DOCX ────────────────────────────────────────────────

def _set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10, color: RGBColor | None = None):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text or '-')
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _add_section_title(doc: DocxDocument, text: str):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x10, 0x2D, 0x45)  # brand-700 do design system
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(4)


def _add_field_row(table, label: str, value: str):
    row = table.add_row().cells
    _set_cell_text(row[0], label.upper(), bold=True, size=8.5, color=RGBColor(0x64, 0x74, 0x8B))
    _set_cell_text(row[1], value or '-', size=10)


def _add_page_number_field(paragraph):
    """Campo 'Página X de Y' — Word não tem essa combinação pronta no
    python-docx, então monta o XML de campo diretamente (padrão documentado
    da lib para isso)."""
    def _field(instr: str):
        run = paragraph.add_run()
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = instr
        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)

    paragraph.add_run('Página ')
    _field('PAGE')
    paragraph.add_run(' de ')
    _field('NUMPAGES')


def gerar_os_docx(data: OSDocumentData) -> bytes:
    doc = DocxDocument()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2)
    section.top_margin = section.bottom_margin = Cm(2)

    # ── Cabeçalho ──────────────────────────────────────────────────────
    header = doc.add_table(rows=1, cols=2)
    header.autofit = True
    logo_cell, dados_cell = header.rows[0].cells
    if _LOGO_PATH.exists():
        try:
            logo_cell.paragraphs[0].add_run().add_picture(str(_LOGO_PATH), width=Cm(4))
        except Exception:
            logger.warning('Falha ao inserir logo no DOCX da OS', exc_info=True)
    dados_par = dados_cell.paragraphs[0]
    dados_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = dados_par.add_run(EMPRESA_RAZAO)
    run.bold = True
    run.font.size = Pt(11)
    for linha in (f'CNPJ {EMPRESA_CNPJ}', EMPRESA_ENDERECO):
        p = dados_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(linha)
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()
    titulo_par = doc.add_paragraph()
    titulo_run = titulo_par.add_run(f'{data.title} — {data.order.number}')
    titulo_run.bold = True
    titulo_run.font.size = Pt(16)
    subtitulo_par = doc.add_paragraph()
    subtitulo_run = subtitulo_par.add_run(
        f'Gerado em {data.generated_at.strftime("%d/%m/%Y %H:%M")} · Status: {data.order.status.value} · Prioridade: {data.order.priority.value}'
    )
    subtitulo_run.font.size = Pt(9)
    subtitulo_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    if data.kind == 'historico_execucao' and data.status_logs is not None:
        _add_section_title(doc, 'Histórico de execução')
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        _set_cell_text(hdr[0], 'Data/Hora', bold=True)
        _set_cell_text(hdr[1], 'Transição', bold=True)
        _set_cell_text(hdr[2], 'Notas', bold=True)
        for entry in data.status_logs:
            row = table.add_row().cells
            _set_cell_text(row[0], entry.when)
            _set_cell_text(row[1], f'{entry.previous} → {entry.new}')
            _set_cell_text(row[2], entry.notes or '-')
    else:
        # ── Dados gerais ───────────────────────────────────────────────
        for titulo, entidade, campos in (
            ('Cliente', data.client, [
                ('Nome/Razão social', getattr(data.client, 'name', None)),
                ('CPF/CNPJ', getattr(data.client, 'cpf_cnpj', None)),
                ('Telefone', getattr(data.client, 'phone', None)),
                ('E-mail', getattr(data.client, 'email', None)),
            ]),
            ('Veículo', data.vehicle, [
                ('Placa', getattr(data.vehicle, 'plate', None)),
                ('Marca/Modelo', ' '.join(filter(None, [getattr(data.vehicle, 'brand', None), getattr(data.vehicle, 'model', None)])) or None),
                ('Chassi', getattr(data.vehicle, 'chassis', None)),
                ('Cor', getattr(data.vehicle, 'color', None)),
            ]),
            ('Rastreador', data.tracker, [
                ('IMEI', getattr(data.tracker, 'imei', None)),
                ('Marca/Modelo', ' '.join(filter(None, [getattr(data.tracker, 'brand', None), getattr(data.tracker, 'model', None)])) or None),
                ('SIM/ICCID', getattr(data.tracker, 'sim_iccid', None) or getattr(data.tracker, 'sim_number', None)),
            ]),
            ('Técnico responsável', data.technician, [
                ('Nome', getattr(data.technician, 'name', None)),
                ('E-mail', getattr(data.technician, 'email', None)),
            ]),
        ):
            if entidade is None:
                continue
            _add_section_title(doc, titulo)
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Light List Accent 1'
            for label, value in campos:
                _add_field_row(table, label, value)

        _add_section_title(doc, 'Agendamento e execução')
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light List Accent 1'
        _add_field_row(table, 'Tipo de serviço', data.order.type.value)
        _add_field_row(table, 'Agendamento', data.order.scheduled_at.strftime('%d/%m/%Y %H:%M') if data.order.scheduled_at else None)
        _add_field_row(table, 'Data de abertura', data.order.created_at.strftime('%d/%m/%Y %H:%M') if data.order.created_at else None)
        _add_field_row(table, 'Data de conclusão', data.order.executed_at.strftime('%d/%m/%Y %H:%M') if data.order.executed_at else None)

        if data.order.problem_description:
            _add_section_title(doc, 'Descrição do problema')
            doc.add_paragraph(data.order.problem_description)

        if data.order.execution_description:
            _add_section_title(doc, 'Serviço executado')
            doc.add_paragraph(data.order.execution_description)

        _add_section_title(doc, 'Checklist técnico')
        if data.checklist:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            _set_cell_text(hdr[0], '', bold=True)
            _set_cell_text(hdr[1], 'Item', bold=True)
            _set_cell_text(hdr[2], 'Observação', bold=True)
            for item in data.checklist:
                row = table.add_row().cells
                # '[X]'/'[ ]' em vez de glifo Unicode: mesmo raciocínio do
                # gerador de PDF (service_order_pdf.py) — garante que
                # feito/não feito sempre fica visualmente distinto,
                # independente da fonte padrão do Word disponível.
                _set_cell_text(row[0], '[X]' if item.done else '[ ]', size=10)
                _set_cell_text(row[1], item.description)
                _set_cell_text(row[2], item.notes or '-')
        else:
            doc.add_paragraph('Nenhum item de checklist registrado.')

        _add_section_title(doc, 'Materiais utilizados')
        if data.materials:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            for i, label in enumerate(('Descrição', 'Qtd.', 'Unid.', 'Preço unit.')):
                _set_cell_text(hdr[i], label, bold=True)
            for m in data.materials:
                row = table.add_row().cells
                _set_cell_text(row[0], m.description)
                _set_cell_text(row[1], m.quantity)
                _set_cell_text(row[2], m.unit or '-')
                _set_cell_text(row[3], m.unit_price or '-')
        else:
            doc.add_paragraph('Nenhum material registrado.')

        if data.order.observations:
            _add_section_title(doc, 'Observações')
            doc.add_paragraph(data.order.observations)

        if data.photos:
            _add_section_title(doc, 'Fotos')
            for photo in data.photos:
                try:
                    doc.add_picture(io.BytesIO(photo.content), width=Cm(8))
                except Exception:
                    logger.warning('Falha ao inserir foto %s no DOCX da OS', photo.file_name, exc_info=True)

        _add_section_title(doc, 'Assinaturas')
        assinatura_table = doc.add_table(rows=2, cols=2)
        for col, (signature_png, label, signed_at) in enumerate((
            (data.technician_signature_png, data.technician.name if data.technician else 'Técnico', data.order.technician_signed_at),
            (data.client_signature_png, data.client.name if data.client else 'Cliente', data.order.client_signed_at),
        )):
            cell = assinatura_table.rows[0].cells[col]
            if signature_png:
                try:
                    cell.paragraphs[0].add_run().add_picture(io.BytesIO(signature_png), width=Cm(6))
                except Exception:
                    logger.warning('Falha ao inserir assinatura no DOCX da OS', exc_info=True)
            legenda_cell = assinatura_table.rows[1].cells[col]
            texto = f'{label}'
            if signed_at:
                texto += f'\nAssinado em {signed_at.strftime("%d/%m/%Y %H:%M")}'
            _set_cell_text(legenda_cell, texto, size=8.5)

    # ── Rodapé (repete em todas as páginas) ──────────────────────────────
    footer_par = section.footer.paragraphs[0]
    footer_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_par.add_run(f'{data.order.number} · ')
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    _add_page_number_field(footer_par)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
